from pathlib import Path
from docx import Document

def main() -> None:
    out = Path(__file__).parent / "sample.docx"
    d = Document()
    d.add_heading("Document Title", level=1)
    d.add_paragraph("Intro paragraph.")
    d.add_heading("Section A", level=2)
    d.add_paragraph("Body of section A.")
    d.add_heading("Section B", level=2)
    d.add_paragraph("Body of section B.")
    d.save(str(out))

if __name__ == "__main__":
    main()
