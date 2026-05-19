from __future__ import annotations

import io

from docx import Document as DocxDocument

from core.exceptions import AppError, ErrorCode
from core.ingestion.parsers import register
from core.ingestion.types import ParsedDocument, ParsedSection


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split(" ", 1)[1])
        except ValueError:
            return None
    return None


class DocxParser:
    kind = "docx"

    def parse_bytes(self, data: bytes, *, source_hint: str | None = None) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(data))
        except Exception as exc:
            raise AppError(
                ErrorCode.INGESTION_PARSE_FAILED, "docx parse failed", detail=str(exc)
            ) from exc

        title: str | None = None
        heading_stack: list[tuple[int, str]] = []
        sections: list[ParsedSection] = []
        buffer: list[str] = []
        ord_ = 0

        def flush() -> None:
            nonlocal buffer, ord_
            body = "\n".join(buffer).strip()
            if not body:
                return
            path = [label for _, label in heading_stack]
            sections.append(
                ParsedSection(text=body, page=None, heading_path=path, ord=ord_)
            )
            ord_ += 1
            buffer = []

        for para in doc.paragraphs:
            level = _heading_level(para.style.name if para.style else None)
            text = para.text.strip()
            if level:
                flush()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))
                if level == 1 and title is None:
                    title = text or None
            else:
                if text:
                    buffer.append(text)
        flush()

        if not sections:
            raise AppError(
                ErrorCode.INGESTION_PARSE_FAILED, "no extractable text in docx"
            )
        return ParsedDocument(title=title or (source_hint or "document.docx"), sections=sections)


register(DocxParser())
